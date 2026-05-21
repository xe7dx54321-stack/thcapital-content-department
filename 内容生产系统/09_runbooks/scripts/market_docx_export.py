#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


KV_RE = re.compile(r"^- `([^`]+)`: ?`?(.*?)`?$")
BULLET_RE = re.compile(r"^- (.+)$")
NUMBER_RE = re.compile(r"^\d+[.)](?:\s+)(.+)$")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export TH Capital market draft pack to formatted DOCX")
    parser.add_argument("--draft-pack-dir", required=True, help="Draft pack directory")
    parser.add_argument("--output-path", default="", help="Output docx path")
    return parser.parse_args()


def clean(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip().strip("`")


def parse_card(path: Path) -> tuple[dict[str, str], dict[str, str]]:
    fields: dict[str, str] = {}
    pack_paths: dict[str, str] = {}
    section = ""
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if line.startswith("## "):
            section = line[3:].strip()
            continue
        match = KV_RE.match(line.strip())
        if not match:
            continue
        key, value = match.groups()
        if section == "Pack Paths":
            pack_paths[key] = clean(value)
        else:
            fields[key] = clean(value)
    return fields, pack_paths


def ensure_styles(document: Document) -> None:
    styles = document.styles

    if "TH Meta" not in styles:
        style = styles.add_style("TH Meta", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor(90, 90, 90)

    if "TH Quote" not in styles:
        style = styles.add_style("TH Quote", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(11.5)
        style.font.italic = False
        style.paragraph_format.left_indent = Inches(0.2)
        style.paragraph_format.right_indent = Inches(0.1)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)

    if "TH Callout" not in styles:
        style = styles.add_style("TH Callout", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = RGBColor(10, 89, 168)


def set_default_font(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11.5)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def add_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_title_page(document: Document, title: str, fields: dict[str, str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.size = Pt(20)

    subtitle = document.add_paragraph(style="TH Meta")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("同行资本市场内容系统｜终稿排版包")

    document.add_paragraph("")

    meta_rows = [
        ("主题", fields.get("draft_key", "n/a")),
        ("核心判断", fields.get("core_judgment", "n/a")),
        ("研究角度", fields.get("approved_angle", "n/a")),
        ("风险提示", fields.get("risk_note", "n/a")),
        ("平台范围", fields.get("requested_platforms", "n/a")),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in meta_rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def render_markdown(document: Document, title: str, content: str) -> None:
    document.add_section(WD_SECTION_START.NEW_PAGE)
    header = document.add_paragraph(style="Heading 1")
    header.add_run(title)

    lines = content.splitlines()
    in_title_options = False
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph(style="Heading 1")
            paragraph.add_run(stripped[2:].strip())
            continue
        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            paragraph = document.add_paragraph(style="Heading 2")
            paragraph.add_run(heading_text)
            in_title_options = heading_text == "标题候选"
            continue
        if stripped.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 3")
            paragraph.add_run(stripped[4:].strip())
            continue
        if stripped.startswith("> "):
            paragraph = document.add_paragraph(style="TH Quote")
            paragraph.add_run(stripped[2:].strip())
            add_shading(paragraph, "EEF5FF")
            continue

        bullet_match = BULLET_RE.match(stripped)
        if bullet_match:
            style = "List Bullet"
            if in_title_options:
                style = "List Bullet 2" if "List Bullet 2" in document.styles else "List Bullet"
            paragraph = document.add_paragraph(style=style)
            paragraph.add_run(bullet_match.group(1))
            continue

        number_match = NUMBER_RE.match(stripped)
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(number_match.group(1))
            continue

        paragraph = document.add_paragraph()
        if stripped.startswith("## 参考来源") or stripped.startswith("## References"):
            paragraph.style = document.styles["Heading 2"]
        paragraph.add_run(stripped)


def read_if_exists(path_str: str) -> str:
    if not path_str or path_str == "n/a":
        return ""
    path = Path(path_str)
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8")


def build_docx(draft_pack_dir: Path, output_path: Path) -> None:
    card_path = draft_pack_dir / "00_draft-pack-card.md"
    fields, pack_paths = parse_card(card_path)

    document = Document()
    ensure_styles(document)
    set_default_font(document)

    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = fields.get("draft_key", draft_pack_dir.name).replace("_", " ")
    add_title_page(document, title, fields)

    platform_order = [
        ("wechat_path", "微信公众号终稿"),
        ("xiaohongshu_path", "小红书终稿"),
        ("zhihu_path", "知乎终稿"),
        ("bilibili_path", "B站专栏终稿"),
        ("toutiao_path", "今日头条终稿"),
        ("baijiahao_path", "百家号终稿"),
        ("x_path", "X Thread 终稿"),
    ]

    for key, label in platform_order:
        content = read_if_exists(pack_paths.get(key, ""))
        if content:
            render_markdown(document, label, content)

    citation = read_if_exists(pack_paths.get("citation_block_path", ""))
    if citation:
        render_markdown(document, "参考来源与引证块", citation)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> None:
    args = parse_args()
    draft_pack_dir = Path(args.draft_pack_dir)
    output_path = Path(args.output_path) if args.output_path else draft_pack_dir / f"{draft_pack_dir.name}__final-content-pack.docx"
    build_docx(draft_pack_dir, output_path)
    print(output_path)


if __name__ == "__main__":
    main()
